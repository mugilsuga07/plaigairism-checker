#!/usr/bin/python3
from django.shortcuts import render, redirect
from django.urls import reverse
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from plagiarismchecker.algorithm.main import tracker,totalPercent, uniquePercent, outputLink, text

# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas

from docx import *
from docx import Document
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 
# import doc2text as dt


# from django.shortcuts import redirect
# from xhtml2pdf import pisa
# from io import BytesIO
# from django.template.loader import get_template
# from django.views import View

# Create your views here.
#home

def home(request):
    return render(request, 'pc/index.html')



# ////////////////////////////////////////////////////////////
# def report(request):
#     # links = list(outputLink.keys())
#     # scores = [i for i in outputLink.values()]
#     if request.POST['q']: 
#         totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(request.POST['q'])
#         print("\n list of links:", links)
#         return render(request, 'pc/report.html', {'tracker' : tracker, 'links' : links, 'Scores':scores, 'text':text,'totalPercent':totalPercent, 'uniquePercent':uniquePercent})
# //////////////////////////
totalPercent = 0
uniquePercent = 0
links = []
scores = [] 
text = str 
tracker = {}
context = {
    'tracker' : tracker, 'links' : links, 'Scores':scores, 'text':text,'totalPercent':totalPercent, 'uniquePercent':uniquePercent
}






def test(request):
    #print("request is welcome test")

    print(request.POST['q'])  
    if request.POST['q']: 
        totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(request.POST['q'])
        uniquePercent = 100 - totalPercent
        totalPercent = round(totalPercent,2)
        uniquePercent = round(uniquePercent,2)
    print("\nOutput..!!!",totalPercent, uniquePercent, links,scores, text, tracker)
    print('\ntext:', text)
    print('\nlinks:', links)
    return render(request, 'pc/index.html',{'totalPercent': totalPercent,'uniquePercent': uniquePercent,'links': links,'scores': scores, 'text' : text, 'tracker':tracker})


def filetest(request):
    value = '' 
    print("GET QUERY STARTED!!!\n")  
    print("Request:\n", request.FILES['docfile'])
    if request.method == 'POST':
        file = request.FILES['docfile']
        # Check the file size
        file = request.FILES.get('docfile')
        if file and file.size > (2 * 1024 * 1024):  # 2MB limit (adjust the limit as needed)
            print("\n...........ERROR:File size limit exceeded (2MB).....................")
            return render(request, 'pc/index.html', {'file_size_error': 'File size limit exceeded (2MB). Upload file size less than 2MB to check for plagiarism.'})

        if str(file).endswith(".txt"):
            value = file.read()
            print('txt value:', value)
            value = value.decode('utf-8')
            print('after decode - txt value:', value)
            
        elif str(file).endswith(".docx"):
            document = Document(file)
            for para in document.paragraphs:
                value += para.text
            print('docx-value:', value)
            if isinstance(value, bytes):
                value = value.decode('utf-8')
                print('after decode - docx value:', value)

            """elif str(file).endswith(".doc"):
            documentdoc = dt.Document(file)
            for para in documentdoc.paragraphs:
                value += para.text
            print('doc-value:', value)
            if isinstance(value, bytes):
                value = value.decode('utf-8')
                print('after decode - doc value:', value)"""
            
            
            
        
        elif str(file).endswith(".pdf"):
            # creating a pdf file object 
            pdfFileObj = file

            # creating a pdf reader object 
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

            # print number of pages in the pdf file
            print("Number of pages:", pdfReader.getNumPages())

            # creating a page object 
            pageObj = pdfReader.getPage(0) 

            # extract text from page
            value = pageObj.extractText()
            print('pdf-value:', value)

            # closing the pdf file object 
            pdfFileObj.close()

        text = value

        # Check word count limit
    
        word_count = len(text.split())
        if word_count > 1500:
            print("\n............ERROR: Word count limit exceeded (1500 words).............")
            return render(request, 'pc/index.html', {'word_count_error': 'Word count limit exceeded (1500 words). Upload file with less than 1500 words to check for plagiarism.'})

    print("\n  TEXT IN FILE INPUT :",text)
    totalPercent, uniquePercent, links, scores, text, tracker= main.findSimilarity(text)
    text = str(value)
    text = text.replace('\r\n', '<br>')
    text = text.replace("\xe2\x80\x99","'")
    print("Output..!!! \n","\n totalPercent, uniquePercent, links, scores, text, tracker\n",totalPercent, uniquePercent, links, scores, text, tracker )
    print("\n list of links:", links)
    print("\n scores:", scores)
    
    return render(request, 'pc/index.html',{'links': links, 'totalPercent': totalPercent,'uniquePercent':uniquePercent, 'text' : text,'scores':scores, 'tracker':tracker})





# .........Extrinsic.................

def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 


#two text compare(Text)
def twofiletest1(request):
    if request.method == 'POST':
        q1 = request.POST.get('q1', '')
        q2 = request.POST.get('q2', '')
        print("Submiited text for 1st and 2nd")
        print("text1 (q1):",q1)
        print("text2 (q2):",q2)
    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        print("\n now go to fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])\n")
        totalPercent, uniquePercent, wordlist1, wordlist2 = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
        
    print("Output>>>!!!! \n plag-percent and unique-percent",totalPercent, uniquePercent,"WORDLIST1",wordlist1,"WORDLIST2",wordlist2)
    return render(request, 'pc/doc_compare.html',{'text1':q1,'text2':q2,'totalPercent': totalPercent,'uniquePercent': uniquePercent,'wordlist1': wordlist1,'wordlist2': wordlist2})
    
# EXTRINSIC - FILES COMPARE

def twofilecompare1(request):
    if request.method == 'POST':
        if 'docfile1' in request.FILES and 'docfile2' in request.FILES:
            q1 = request.FILES['docfile1']
            q2 = request.FILES['docfile2']
            print('\nFile 1 (docfile1):', q1.name)
            print('\nFile 2 (docfile2):', q2.name)

            wordlist1 = []
            wordlist2 = []
            result = {}

            if q1.name.endswith(".txt") and q2.name.endswith(".txt"):
                q1 = q1.read().decode('utf-8')
                print('q1 after decode -', q1)

                q2 = q2.read().decode('utf-8')
                print('\nq2 after decode -', q2)

            elif q1.name.endswith(".docx") and q2.name.endswith(".docx"):
                document = Document(q1)
                q1 = ""
                for para in document.paragraphs:
                    q1 += para.text

                document = Document(q2)
                q2 = ""
                for para in document.paragraphs:
                    q2 += para.text

            print("file1 (q1):", q1)
            print('file2 (q2):', q2)
            
            totalPercent, uniquePercent, wordlist1, wordlist2 = fileSimilarity.findFileSimilarity(q1, q2)
            print("Output result:! \n", result)
            
            return render(request, 'pc/doc_compare.html', {
                'text1': q1,'text2': q2,'totalPercent': totalPercent,'uniquePercent': uniquePercent,
                'wordlist1': wordlist1,'wordlist2': wordlist2
            })
    
    return render(request, 'pc/doc_compare.html')


# .................
# def fileCompare(request):
#     return render(request, 'pc/doc_compare.html') 

# def twofiletest1(request):
#     print("Submitted text for 1st and 2nd")
#     print(request.POST['q1'])
#     print(request.POST['q2'])

#     if request.POST['q1'] != '' and request.POST['q2'] != '': 
#         print("Got both the texts")
#         result = fileSimilarity.findFileSimilarity(request.POST['q1'], request.POST['q2'])
#         result = round(result, 2)
#         print("Output>>>!!!!", result)
#         return render(request, 'pc/doc_compare.html', {'result': result})

# def twofilecompare1(request):
#     value1 = ''
#     value2 = ''
#     value3 = ''
#     inputQuery = ''
#     database = ''
#     matchPercentage = 0
#     uniquePercentage = 0


#     if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
#         value1 = str(request.FILES['docfile1'].read())
#         value2 = str(request.FILES['docfile2'].read())

#     elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
#         document = Document(request.FILES['docfile1'])
#         for para in document.paragraphs:
#             value1 += para.text
#         document = Document(request.FILES['docfile2'])
#         for para in document.paragraphs:
#             value2 += para.text

#     elif (str(request.FILES['docfile1'])).endswith(".doc") and (str(request.FILES['docfile2'])).endswith(".doc"):
#         document = Document(request.FILES['docfile1'])
#         for para in document.paragraphs:
#             value1 += para.text
#         document = Document(request.FILES['docfile2'])
#         for para in document.paragraphs:
#             value2 += para.text

#     inputQuery, database = fileSimilarity.findFileSimilarity(inputQuery, database)
#     print("Output! \n", inputQuery, database)
#     return render(request, 'pc/doc_compare.html', {'matchPercentage':matchPercentage, 'uniquePercentage': uniquePercentage
# })
